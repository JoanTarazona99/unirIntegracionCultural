#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from pathlib import Path
import shutil

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(r"c:\xampp\htdocs\proyectos\unirIntegracionCultural")
DOCX = ROOT / "курсовая_090403_Тарасона.docx"
BACKUP = ROOT / "курсовая_090403_Тарасона_BACKUP_RESTRUCTURE.docx"
FIG_DIR = ROOT / "data" / "eval" / "figures"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def make_grounding_figure(out_path: Path):
    fig, ax = plt.subplots(figsize=(8.6, 6.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 12)
    ax.axis("off")

    def box(x, y, w, h, txt, fc="#eef4ff", ec="#2c6fbb", fs=9):
        b = FancyBboxPatch((x - w / 2, y - h / 2), w, h,
                           boxstyle="round,pad=0.08,rounding_size=0.12",
                           fc=fc, ec=ec, lw=1.4)
        ax.add_patch(b)
        ax.text(x, y, txt, ha="center", va="center", fontsize=fs, wrap=True)

    def arr(x1, y1, x2, y2):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>",
                                     mutation_scale=14, lw=1.2, color="#333"))

    ax.text(5, 11.5, "Рисунок: Конвейер grounding и политики воздержания", ha="center", fontsize=11, fontweight="bold")

    box(5, 10.3, 6.0, 1.1, "Сгенерированный ответ LLM (Qwen2.5)", fc="#f3e8ff", ec="#7e57c2")
    box(5, 8.7, 6.8, 1.2, "analyze_grounding_improved():\nлексическое совпадение + hard entities + эвристики", fc="#fff3e0", ec="#ef6c00")
    box(5, 7.0, 6.8, 1.0, "Классификация: HIGH (>=0.75), MEDIUM (0.4-0.75), LOW (<0.4)", fc="#e8f5e9", ec="#2e7d32")

    box(2.7, 5.2, 3.7, 1.0, "Тема чувствительная?\n(visa/registration/fees)", fc="#ffebee", ec="#c62828")
    box(7.3, 5.2, 3.7, 1.0, "enforce_grounding_improved()\n+ citation_guard", fc="#e1f5fe", ec="#1565c0")

    box(2.7, 3.3, 3.9, 1.0, "Порог выше\n(строгий режим)", fc="#ffebee", ec="#c62828")
    box(7.3, 3.3, 3.9, 1.0, "Стандартный порог\n(базовый режим)", fc="#e8f5e9", ec="#2e7d32")

    box(5, 1.5, 8.2, 1.2, "Выход: ответ с цитированием источников\nили контролируемое воздержание с направлением к официальным каналам",
        fc="#f5f5f5", ec="#424242")

    arr(5, 9.75, 5, 9.3)
    arr(5, 8.1, 5, 7.55)
    arr(5, 6.5, 2.9, 5.7)
    arr(5, 6.5, 7.1, 5.7)
    arr(2.7, 4.7, 2.7, 3.8)
    arr(7.3, 4.7, 7.3, 3.8)
    arr(2.7, 2.8, 4.3, 2.0)
    arr(7.3, 2.8, 5.7, 2.0)

    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def make_kb_refresh_figure(out_path: Path):
    fig, ax = plt.subplots(figsize=(9.2, 6.3))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 10)
    ax.axis("off")

    def box(x, y, w, h, txt, fc="#eef4ff", ec="#2c6fbb", fs=9):
        b = FancyBboxPatch((x - w / 2, y - h / 2), w, h,
                           boxstyle="round,pad=0.08,rounding_size=0.12",
                           fc=fc, ec=ec, lw=1.4)
        ax.add_patch(b)
        ax.text(x, y, txt, ha="center", va="center", fontsize=fs, wrap=True)

    def arr(x1, y1, x2, y2):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>",
                                     mutation_scale=14, lw=1.2, color="#333"))

    ax.text(6, 9.5, "Рисунок: Инкрементальное обновление базы знаний", ha="center", fontsize=11, fontweight="bold")

    box(2.2, 7.8, 3.6, 1.2, "Event-driven trigger\n(изменение источника)", fc="#e8f5e9", ec="#2e7d32")
    box(6.0, 7.8, 3.6, 1.2, "Scheduler refresh\n(critical/faq/stable)", fc="#e1f5fe", ec="#1565c0")
    box(9.8, 7.8, 3.6, 1.2, "Leader lock / distributed lock\n(анти-конкуренция)", fc="#fff3e0", ec="#ef6c00")

    box(6.0, 5.9, 6.8, 1.2, "kb_refresh: fetch -> fingerprint -> compare -> version_id", fc="#f3e8ff", ec="#7e57c2")
    box(6.0, 4.2, 6.8, 1.2, "Исключение stale/inactive + инкрементальная переиндексация", fc="#ffebee", ec="#c62828")
    box(6.0, 2.5, 6.8, 1.2, "upsert_refreshed_section + source_registry + kb_versions", fc="#f5f5f5", ec="#424242")
    box(6.0, 0.9, 7.4, 1.2, "Результат: консистентная KB для RAG + журнал обновлений + rollback-ready", fc="#e8f5e9", ec="#2e7d32")

    arr(2.2, 7.2, 4.4, 6.3)
    arr(6.0, 7.2, 6.0, 6.6)
    arr(9.8, 7.2, 7.6, 6.3)
    arr(6.0, 5.3, 6.0, 4.8)
    arr(6.0, 3.6, 6.0, 3.1)
    arr(6.0, 1.9, 6.0, 1.4)

    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    grounding_fig = FIG_DIR / "fig_grounding_pipeline.png"
    refresh_fig = FIG_DIR / "fig_kb_refresh_cycle.png"
    make_grounding_figure(grounding_fig)
    make_kb_refresh_figure(refresh_fig)

    doc = Document(DOCX)

    # 1) Reestructurar capítulos principales
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "Ознакомление с системой":
            p.text = "ГЛАВА 1. АНАЛИЗ ПРОБЛЕМЫ И ПОСТАНОВКА ЗАДАЧИ"
            p.style = doc.styles["Heading 1"]
        elif t == "Проектирование интеллектуального ассистента культурной интеграции":
            p.text = "ГЛАВА 2. ПРОЕКТИРОВАНИЕ И АРХИТЕКТУРА СИСТЕМЫ"
            p.style = doc.styles["Heading 1"]
        elif t == "Реализация и тестирование прототипа":
            p.text = "ГЛАВА 3. РЕАЛИЗАЦИЯ СИСТЕМЫ"
            p.style = doc.styles["Heading 1"]
        elif t == "Тестирование и оценка результатов":
            p.text = "ГЛАВА 4. ОЦЕНКА И ВАЛИДАЦИЯ"
            p.style = doc.styles["Heading 1"]

    # 2) Усиление секции достоверности (citation_guard + abstention)
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "Обеспечение достоверности ответов и контроль галлюцинаций":
            target = p
            break

    if target is not None:
        p1 = insert_paragraph_after(
            target,
            "В текущей реализации контроль достоверности выполняется каскадно: сначала вычисляется показатель grounding/faithfulness, затем применяется модуль citation_guard, который принудительно проверяет наличие подтверждающих источников. Для чувствительных доменов (виза, регистрация, административные процедуры) применяется более строгий порог принятия ответа.",
            style=doc.styles["Normal"],
        )
        p2 = insert_paragraph_after(
            p1,
            "Если уровень обоснованности ниже порога или источники недостаточны, система активирует политику воздержания: вместо гипотетического ответа выдаётся безопасное сообщение с направлением пользователя к официальным каналам (университет, миграционные службы, МФЦ). Такой режим уменьшает риск распространения недостоверных сведений.",
            style=doc.styles["Normal"],
        )

        pimg = insert_paragraph_after(p2, style=doc.styles["Normal"])
        pimg.alignment = 1
        pimg.add_run().add_picture(str(grounding_fig), width=Inches(5.8))
        pcap = insert_paragraph_after(
            pimg,
            "Рисунок 4 – Конвейер grounding, citation_guard и политики воздержания",
            style=doc.styles["Normal"],
        )
        pcap.alignment = 1

    # 3) Добавить главу 5 перед заключением
    conclusion = None
    for p in doc.paragraphs:
        if p.text.strip() == "Заключение":
            conclusion = p
            break

    if conclusion is not None:
        h5 = conclusion.insert_paragraph_before("ГЛАВА 5. ЭКСПЛУАТАЦИЯ, НАДЁЖНОСТЬ И ЭВОЛЮЦИЯ БАЗЫ ЗНАНИЙ")
        h5.style = doc.styles["Heading 1"]

        s51 = conclusion.insert_paragraph_before("Обновление базы знаний и управление версиями")
        s51.style = doc.styles["Heading 3"]

        b1 = conclusion.insert_paragraph_before(
            "Операционный контур системы построен на инкрементальном обновлении базы знаний: для каждого источника вычисляется fingerprint, формируется version_id, а затем выполняется точечное обновление только изменившихся секций. Такой подход уменьшает вычислительную нагрузку и снижает риск нарушения консистентности индекса."
        )
        b1.style = doc.styles["Normal"]

        b2 = conclusion.insert_paragraph_before(
            "В процессе обновления поддерживается политика исключения неактивного и устаревшего контента (inactive/stale), а также сохраняется история версий для трассируемости и возможности rollback. В результате RAG-модуль работает с актуальным и проверяемым набором данных."
        )
        b2.style = doc.styles["Normal"]

        bimg = conclusion.insert_paragraph_before("")
        bimg.style = doc.styles["Normal"]
        bimg.alignment = 1
        bimg.add_run().add_picture(str(refresh_fig), width=Inches(6.0))

        bcap = conclusion.insert_paragraph_before("Рисунок 5 – Инкрементальное обновление базы знаний: event-driven + scheduled refresh")
        bcap.style = doc.styles["Normal"]
        bcap.alignment = 1

        s52 = conclusion.insert_paragraph_before("Конкурентный доступ и multi-worker устойчивость")
        s52.style = doc.styles["Heading 3"]

        b3 = conclusion.insert_paragraph_before(
            "Для предотвращения гонок при обновлении источников используется механизм блокировок (state lock) и лидерный контур запуска scheduler-задач. В многоинстансных конфигурациях может применяться distributed lock, что исключает параллельное выполнение критических операций несколькими worker-процессами."
        )
        b3.style = doc.styles["Normal"]

        s53 = conclusion.insert_paragraph_before("Практическая валидация и эксплуатационные метрики")
        s53.style = doc.styles["Heading 3"]

        b4 = conclusion.insert_paragraph_before(
            "Эксплуатационная валидация выполняется сочетанием автоматических тестов и ручных сценариев refresh-пайплайна: проверяются корректность детекции изменений, стабильность переиндексации, соблюдение политики воздержания и целостность ссылок на источники. Эти процедуры позволяют использовать систему в режиме контролируемой эволюции без деградации качества ответов."
        )
        b4.style = doc.styles["Normal"]

    # 4) Актуализировать подписи старых графиков оценки после добавления новых рисунков
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Рисунок 4 – Сравнение метрик качества поиска"):
            p.text = "Рисунок 6 – Сравнение метрик качества поиска для различных режимов работы модуля"
        elif t.startswith("Рисунок 5 – Качество поиска по тематическим категориям"):
            p.text = "Рисунок 7 – Качество поиска по тематическим категориям (nDCG@5)"

    # 5) Добавить переход в главе 3 к новой главе 4
    for p in doc.paragraphs:
        if p.text.strip() == "Реализация пользовательских интерфейсов":
            trans = insert_paragraph_after(
                p,
                "Представленные компоненты реализации образуют технологическую основу прототипа. Далее в отдельной главе приводится формальная валидация качества поиска, устойчивости и достоверности ответов на основе количественных метрик и экспериментальных сценариев.",
                style=doc.styles["Normal"],
            )
            break

    doc.save(DOCX)
    print("DONE: document restructured and figures added")


if __name__ == "__main__":
    main()
