#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para reemplazar figuras en el documento DOCX.
Reemplaza Figura 1-3 con las nuevas versiones profesionales.
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = '/c/xampp/htdocs/proyectos/unirIntegracionCultural/курсовая_090403_Тарасона.docx'
figures_dir = '/c/xampp/htdocs/proyectos/unirIntegracionCultural/data/eval/figures'

# Mapping de figuras a reemplazar: (número de figura, descripción, nueva imagen)
FIGURES_TO_UPDATE = [
    (1, 'fig_arquitectura.png', '§2.2 Arquitectura general del sistema'),
    (2, 'fig_rag_detalle.png', '§2.4 Esquema del módulo RAG'),
    (3, 'fig_voz.png', '§3.3 Pipeline de voz'),
]

def find_and_replace_figures():
    """Encuentra referencias a figuras y las reemplaza con las nuevas imágenes."""
    doc = Document(doc_path)
    
    # Primero, encontrar dónde están las referencias a "Рисунок 1", "Рисунок 2", etc.
    replaced_count = 0
    
    for para_idx, para in enumerate(doc.paragraphs):
        # Buscar párrafos que mencionen "Рисунок X"
        for fig_num, img_file, description in FIGURES_TO_UPDATE:
            pattern = f"Рисунок {fig_num}"
            
            if pattern in para.text and para.text.strip().startswith(pattern):
                # Este es el párrafo que describe la figura
                # Buscar la siguiente imagen para reemplazarla
                
                # Escanear hacia adelante para encontrar la imagen
                img_path = os.path.join(figures_dir, img_file)
                
                if os.path.exists(img_path):
                    print(f"✓ Encontrada referencia: {pattern} en párrafo {para_idx}")
                    print(f"  Descripción: {description}")
                    print(f"  Reemplazando con: {img_file}")
                    
                    # Buscar la imagen en los párrafos siguientes (típicamente dentro de los próximos 2-3 párrafos)
                    for next_idx in range(para_idx + 1, min(para_idx + 5, len(doc.paragraphs))):
                        next_para = doc.paragraphs[next_idx]
                        
                        # Buscar runs con imágenes
                        for run in next_para.runs:
                            # Buscar shapes (imágenes embedidas)
                            if hasattr(run, '_element'):
                                inline_shapes = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                                if inline_shapes:
                                    # Encontrada una imagen, reemplazarla
                                    # Limpiar la imagen existente y agregar la nueva
                                    run.clear()
                                    run.add_picture(img_path, width=Inches(5.5))
                                    replaced_count += 1
                                    print(f"  ✓ Imagen reemplazada en run {para_idx}")
                                    break
    
    # Si no se encontraron mediante ese método, intenta agregar después de los párrafos de descripción
    if replaced_count == 0:
        print("\n⚠️  Método directo no encontró imágenes. Buscando alternativa...")
        # Buscar manualmente todos los párrafos que mencionen "Рисунок"
        for para_idx, para in enumerate(doc.paragraphs):
            for fig_num, img_file, description in FIGURES_TO_UPDATE:
                if f"Рисунок {fig_num}" in para.text:
                    print(f"✓ {description}")
                    # Simplemente agregar la imagen después del párrafo de descripción
                    # Crear un nuevo párrafo e insertar la imagen
                    new_para = para._element.addnext(doc.add_paragraph()._element)
                    new_p_obj = doc.paragraphs[-1]  # El último párrafo agregado
                    new_p_obj.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    img_path = os.path.join(figures_dir, img_file)
                    if os.path.exists(img_path):
                        new_p_obj.add_run().add_picture(img_path, width=Inches(5.5))
                        replaced_count += 1
                        print(f"  ✓ Agregada nueva imagen: {img_file}")
    
    # Guardar documento actualizado
    doc.save(doc_path)
    print(f"\n{'='*70}")
    print(f"✓ Documento actualizado: {doc_path}")
    print(f"  Figuras actualizadas: {replaced_count}")
    print(f"{'='*70}")

if __name__ == '__main__':
    print(f"Reemplazando figuras en documento...")
    print(f"Documento: {doc_path}")
    print(f"Directorio de figuras: {figures_dir}")
    print()
    find_and_replace_figures()
