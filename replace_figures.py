#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Reemplazar figuras 1-3 en el DOCX con las nuevas versiones profesionales.
Acceso directo a elementos XML del documento.
"""
import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt

doc_path = r'c:\xampp\htdocs\proyectos\unirIntegracionCultural\курсовая_090403_Тарасона.docx'
figures_dir = r'c:\xampp\htdocs\proyectos\unirIntegracionCultural\data\eval\figures'

# Índices de párrafos donde están las referencias + nueva imagen a usar
REPLACEMENTS = [
    (159, 'fig_arquitectura.png', 'Arquitectura del Sistema'),
    (210, 'fig_rag_detalle.png', 'Esquema RAG'),
    (325, 'fig_voz.png', 'Pipeline de Voz'),
]

def replace_figure_after_paragraph(doc, para_idx, img_file, description):
    """Reemplaza o agrega imagen después del párrafo indicado."""
    img_path = os.path.join(figures_dir, img_file)
    
    if not os.path.exists(img_path):
        print(f"❌ Archivo no encontrado: {img_path}")
        return False
    
    para = doc.paragraphs[para_idx]
    print(f"\n▸ Procesando párrafo [{para_idx}]: {description}")
    print(f"  Imagen: {img_file}")
    
    # Buscar la imagen en los siguientes párrafos (típicamente 1-2 párrafos después)
    found_image = False
    for next_idx in range(para_idx + 1, min(para_idx + 5, len(doc.paragraphs))):
        next_para = doc.paragraphs[next_idx]
        
        # Buscar runs que contengan imágenes
        for run_idx, run in enumerate(next_para.runs):
            # Verificar si hay dibujos (imágenes)
            drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawings:
                print(f"  ▹ Imagen encontrada en run {run_idx} del párrafo [{next_idx}]")
                # Reemplazar la imagen
                # Primero, limpiar el run
                run._element.clear_content()
                
                # Agregar la nueva imagen
                run.add_picture(img_path, width=Inches(5.5))
                found_image = True
                print(f"  ✓ Imagen reemplazada")
                break
        
        if found_image:
            break
    
    if not found_image:
        print(f"  ⚠️  No se encontró imagen existente")
        # Crear un nuevo párrafo con la imagen
        new_para = para._element.addnext(doc.add_paragraph()._element)
        # Actualizar índices... es complicado, así que mejor agregar al final del párrafo actual
        new_run = para.add_run()
        new_run.add_picture(img_path, width=Inches(5.5))
        print(f"  ✓ Nueva imagen agregada al párrafo")
    
    return True

def main():
    print("="*70)
    print("REEMPLAZANDO FIGURAS EN DOCUMENTO")
    print("="*70)
    
    doc = Document(doc_path)
    print(f"\n Documento: {doc_path}")
    print(f"Directorio de figuras: {figures_dir}")
    print(f"Total de párrafos: {len(doc.paragraphs)}")
    
    success_count = 0
    for para_idx, img_file, desc in REPLACEMENTS:
        if replace_figure_after_paragraph(doc, para_idx, img_file, desc):
            success_count += 1
    
    # Guardar
    doc.save(doc_path)
    
    print("\n" + "="*70)
    print(f"✓ COMPLETADO: {success_count} figuras reemplazadas")
    print(f"✓ Documento guardado: {doc_path}")
    print("="*70)

if __name__ == '__main__':
    main()
